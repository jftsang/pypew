import datetime as dt
import os
import re
import typing
from abc import ABC, abstractmethod
from functools import lru_cache
from pathlib import Path
from typing import List, Optional

import jinja2
import yaml
from attr import define, field
from dateutil.easter import easter
from docx import Document
from docxtpl import DocxTemplate, RichText

from models_base import get

if typing.TYPE_CHECKING:
    from forms import PewSheetForm

from utils import get_neh_df, advent, closest_sunday_to, NoPandasError, logger

feasts_fields = ['name', 'month', 'day', 'coeaster', 'coadvent',
                 'introit', 'collect', 'epistle_ref', 'epistle',
                 'gat', 'gradual', 'alleluia', 'tract', 'gospel_ref',
                 'gospel', 'offertory', 'communion']
DATA_DIR = Path(os.path.dirname(__file__)) / 'data' / 'feasts'
PEW_SHEET_TEMPLATE = os.path.join('templates', 'pewSheetTemplate.docx')


def _none2datemax(d: Optional[dt.date]) -> dt.date:
    """Put unspecified dates at the end of the list."""
    if d is None:
        return dt.date.max
    return d


@define
class Feast:
    @classmethod
    def from_yaml(cls, slug):
        return _feast_from_yaml(slug)

    @classmethod
    def all(cls):
        with open(DATA_DIR / '_list.txt') as f:
            slugs = [x.strip() for x in f]
            return [cls.from_yaml(slug) for slug in slugs]

    @classmethod
    def upcoming(cls, date: Optional[dt.date] = None) -> List['Feast']:
        if date is None:
            date = dt.date.today()

        return sorted(Feast.all(),
                      key=lambda f: _none2datemax(f.get_next_date(date)))

    @classmethod
    def next(cls, date: Optional[dt.date] = None) -> 'Feast':
        if date is None:
            date = dt.date.today()

        return min(Feast.all(),
                   key=lambda f: _none2datemax(f.get_next_date(date)))

    @classmethod
    def get(cls, **kwargs):
        return get(cls.all(), **kwargs)

    slug: str = field()
    name: str = field()

    # Specified for the fixed holy days, None for the movable feasts.
    # TODO - what about Remembrance Sunday and Advent Sunday? Not fixed
    #  days but also not comoving with Easter. As a hack go with 11 Nov
    #  and 30 Nov respectively but the exact dates are
    month: Optional[int] = field(default=None)
    day: Optional[int] = field(default=None)

    # For the feasts synced with Easter, the number of days since Easter
    coeaster: Optional[int] = field(default=None)
    coadvent: Optional[int] = field(default=None)

    introit: Optional[str] = field(default=None)
    collect: Optional[str] = field(default=None)
    epistle_ref: Optional[str] = field(default=None)
    epistle: Optional[str] = field(default=None)
    gat: str = field(default='')
    gradual: Optional[str] = field(default=None)
    alleluia: Optional[str] = field(default=None)
    tract: Optional[str] = field(default=None)
    gospel_ref: Optional[str] = field(default=None)
    gospel: Optional[str] = field(default=None)
    offertory: Optional[str] = field(default=None)
    communion: Optional[str] = field(default=None)

    def get_date(self, year=None) -> Optional[dt.date]:
        if year is None:
            year = dt.datetime.now().year

        if self.month is not None and self.day is not None:
            # TODO Check this definition
            if self.name == 'Remembrance Sunday':
                return closest_sunday_to(dt.date(year, self.month, self.day))

            return dt.date(year, self.month, self.day)

        assert not (self.coeaster is not None and self.coadvent is not None)

        if self.coeaster is not None:
            return easter(year) + dt.timedelta(days=self.coeaster)

        if self.coadvent is not None:
            return advent(year) + dt.timedelta(days=self.coadvent)

        return None

    @property
    def date(self):
        """The date of the feast in the present year."""
        return self.get_date()

    def get_next_date(self, d: Optional[dt.date] = None) -> Optional[dt.date]:
        """Returns the next occurrence of this feast from the specified
        date, which may be in the next calendar year.
        """
        if d is None:
            d = dt.date.today()

        next_occurrence = self.get_date(year=d.year)
        if next_occurrence is None:
            return None
        if (next_occurrence - d).days < 0:
            next_occurrence = self.get_date(year=d.year + 1)

        return next_occurrence

    @property
    def next_date(self):
        """The next occurrence of the feast, which may be in the next
        calendar year.
        """
        return self.get_next_date()

    def create_docx(self, path):
        document = Document()
        document.add_heading(self.name, 0)
        document.save(path)


@define
class DateRule:
    # Specified for the fixed holy days, None for the movable feasts.
    # TODO - what about Remembrance Sunday and Advent Sunday? Not fixed
    #  days but also not comoving with Easter. As a hack go with 11 Nov
    #  and 30 Nov respectively but the exact dates are
    month: Optional[int] = field(default=None)
    day: Optional[int] = field(default=None)

    # For the feasts synced with Easter, the number of days since Easter
    coeaster: Optional[int] = field(default=None)
    coadvent: Optional[int] = field(default=None)


class PewSheetItem(ABC):
    @abstractmethod
    def as_richtext(self) -> RichText:
        raise NotImplementedError


class CharacterStyles:
    title = {'font': 'Merriweather', 'size': 20, 'bold': True}
    subtitle = {'font': 'Raleway', 'size': 18, 'italic': True}
    paragraph = {'font': 'Cambria', 'size': 20}


@define
class Music:
    title: str = field()
    category: str = field()  # Anthem or Hymn or Plainsong
    composer: Optional[str] = field()
    lyrics: Optional[str] = field()
    ref: Optional[str] = field()
    translation: Optional[str] = field()

    @classmethod
    def neh_hymns(cls) -> List['Music']:
        try:
            records = get_neh_df().itertuples()
        except NoPandasError as exc:
            logger.warning(exc)
            return []

        def nehref2num(nehref: str) -> typing.Tuple[int, str]:
            m = re.match(r"NEH: (\d+)([a-z]?)", nehref)
            assert m is not None
            num, suffix = m.groups()
            return int(num), suffix

        hymns = [
            Music(
                title=record.firstLine,
                category='Hymn',
                composer=None,
                lyrics=None,
                ref=f'NEH: {record.number}',
                translation=f'Words/translation available at NEH: {record.number}, {record.firstLine}'
            ) for record in records
        ]
        hymns.sort(key=lambda m: nehref2num(m.ref or ""))
        return hymns

    @classmethod
    def get_neh_hymn_by_ref(cls, ref: str) -> Optional['Music']:
        try:
            return next(filter(lambda h: h.ref == ref, cls.neh_hymns()))
        except StopIteration:
            return None

    def __str__(self):
        if self.category == 'Hymn':
            return f'{self.ref}, {self.title}'
        return super().__str__()

    def as_richtext(self) -> RichText:
        if self.category != 'Hymn':
            return RichText(super().__str__())

        assert self.ref is not None

        rt = RichText()
        rt.add(self.ref + ', ', **CharacterStyles.paragraph)
        rt.add(self.title, **CharacterStyles.paragraph, italic=True)
        return rt


@define
class ServiceItem(PewSheetItem):
    title: str = field(default='')
    paragraphs: List[typing.Any] = field(factory=list)
    subtitle: Optional[str] = field(default=None)

    def as_richtext(self) -> RichText:

        rt = RichText()
        rt.add(self.title, **CharacterStyles.title)
        if self.subtitle is not None:
            rt.add('\a')
            rt.add(self.subtitle, **CharacterStyles.subtitle)

        for paragraph in self.paragraphs:
            rt.add('\a')
            rt.add(paragraph, **CharacterStyles.paragraph)

        return rt


@define
class CollectItem(PewSheetItem):
    collects: List[str] = field(factory=list)

    @property
    def title(self):
        return 'Collect' if len(self.collects) == 1 else 'Collects'

    @property
    def subtitle(self):
        return None

    @property
    def paragraphs(self):
        return self.collects

    def as_richtext(self) -> RichText:
        rt = RichText()
        rt.add(self.title, **CharacterStyles.title)
        for collect in self.collects:
            if collect.endswith('Amen.'):
                rt.add('\a' + collect[:-5], **CharacterStyles.paragraph)
                rt.add(collect[-5:], **CharacterStyles.paragraph, bold=True)
            else:
                rt.add('\a' + collect, **CharacterStyles.paragraph)

        return rt




@define
class MusicItem(PewSheetItem):
    title: str = field()
    music: Music = field()

    @property
    def subtitle(self):
        return None

    @property
    def paragraphs(self):
        return [str(self.music)]

    def as_richtext(self) -> RichText:
        rt = RichText()
        rt.add(self.title, font='Merriweather', size=20, bold=True)
        rt.add('\a')
        rt.add(self.music.as_richtext())
        return rt


@define
class Service:
    # Mandatory fields first, then fields with default values.
    title: str = field()
    date: dt.date = field()
    primary_feast: Feast = field()
    time: dt.time = field(default=dt.time(11, 0))
    secondary_feasts: [Feast] = field(factory=list)
    celebrant: str = field(default='')
    preacher: str = field(default='')
    introit_hymn: Optional[Music] = field(default=None)
    offertory_hymn: Optional[Music] = field(default=None)
    recessional_hymn: Optional[Music] = field(default=None)
    anthem: Optional[Music] = field(default=None)
    service_type: str = field(default='Sung Mass')

    # One can't call methods in jinja2 templates, so one must provide
    # everything as member properties instead.

    @property
    def collects(self) -> List[str]:
        out = []
        if self.primary_feast.collect:
            out.append(self.primary_feast.collect)
        for sf in self.secondary_feasts:
            if sf.collect:
                out.append(sf.collect)

        # Collects for Advent I and Ash Wednesday are repeated
        # throughout Advent and Lent respectively.
        advent1 = Feast.get(name='Advent I')
        ash_wednesday = Feast.get(name='Ash Wednesday')

        if 'Advent' in self.primary_feast.name and self.primary_feast != advent1:
            out.append(advent1.collect)

        if 'Lent' in self.primary_feast.name:
            out.append(ash_wednesday.collect)

        return out

    # TODO primary or secondary?
    @property
    def introit_proper(self) -> Optional[str]:
        return self.primary_feast.introit

    @property
    def gat(self) -> str:
        assert self.primary_feast.gat is not None
        return self.primary_feast.gat

    @property
    def gat_propers(self) -> List[str]:
        propers = []
        if 'Gradual' in self.primary_feast.gat:
            assert self.primary_feast.gradual is not None
            propers.append(self.primary_feast.gradual)
        if 'Alleluia' in self.primary_feast.gat:
            assert self.primary_feast.alleluia is not None
            propers.append(self.primary_feast.alleluia)
        if 'Tract' in self.primary_feast.gat:
            assert self.primary_feast.tract is not None
            propers.append(self.primary_feast.tract)
        return propers

    @property
    def offertory_proper(self) -> Optional[str]:
        return self.primary_feast.offertory

    @property
    def communion_proper(self) -> Optional[str]:
        return self.primary_feast.communion

    @property
    def epistle_ref(self) -> Optional[str]:
        return self.primary_feast.epistle_ref

    @property
    def epistle(self) -> Optional[str]:
        return self.primary_feast.epistle

    @property
    def gospel_ref(self) -> Optional[str]:
        return self.primary_feast.gospel_ref

    @property
    def gospel(self) -> Optional[str]:
        return self.primary_feast.gospel

    @property
    def items(self) -> List[PewSheetItem]:
        items: List[PewSheetItem] = []
        if self.introit_hymn:
            items.append(
                MusicItem('Introit Hymn', self.introit_hymn)
            )
        items.append(ServiceItem('Introit Proper', [self.introit_proper]))

        collects = CollectItem(self.collects)
        items.append(collects)

        items.append(ServiceItem('Epistle', [self.epistle], self.epistle_ref))
        items.append(ServiceItem(self.gat, self.gat_propers))
        items.append(ServiceItem('Gospel', [self.gospel], self.gospel_ref))

        items.append(ServiceItem('Offertory Proper', [self.offertory_proper]))
        if self.offertory_hymn:
            items.append(MusicItem('Offertory Hymn', self.offertory_hymn))

        items.append(ServiceItem('Communion Proper', [self.communion_proper]))

        if self.anthem:
            items.append(
                ServiceItem('Anthem', [self.anthem.lyrics, self.anthem.translation],
                            f'{self.anthem.title}. {self.anthem.composer}'))

        if self.recessional_hymn:
            items.append(
                MusicItem('Recessional Hymn', self.recessional_hymn))

        return items

    @classmethod
    def from_form(cls, form: 'PewSheetForm') -> 'Service':
        primary_feast = Feast.get(slug=form.primary_feast.data)
        if form.secondary_feasts.data:
            secondary_feasts = [Feast.get(slug=slug)
                                for slug in form.secondary_feasts.data
                                if slug]
        else:
            secondary_feasts = []

        if form.anthem_title.data or form.anthem_composer.data or form.anthem_lyrics.data or form.anthem_translation.data:
            anthem = Music(
                title=form.anthem_title.data,
                composer=form.anthem_composer.data,
                lyrics=form.anthem_lyrics.data,
                category='Anthem',
                ref=None,
                translation=form.anthem_translation.data
            )
        else:
            anthem = None

        return Service(
            title=form.title.data,
            date=form.date.data,
            time=form.time.data,
            celebrant=form.celebrant.data,
            preacher=form.preacher.data,
            primary_feast=primary_feast,
            secondary_feasts=secondary_feasts,
            introit_hymn=Music.get_neh_hymn_by_ref(form.introit_hymn.data),
            offertory_hymn=Music.get_neh_hymn_by_ref(form.offertory_hymn.data),
            recessional_hymn=Music.get_neh_hymn_by_ref(
                form.recessional_hymn.data
            ),
            anthem=anthem,
        )

    def create_docx(self, path):
        doc = DocxTemplate(PEW_SHEET_TEMPLATE)
        jinja_env = jinja2.Environment(autoescape=True)
        jinja_env.globals['len'] = len

        # local import to avoid circular import
        from filters import filters_context

        jinja_env.filters.update(filters_context)
        doc.render({'service': self}, jinja_env)
        doc.save(path)


@lru_cache()
def _feast_from_yaml(slug: str) -> Feast:
    with open((DATA_DIR / slug).with_suffix('.yaml')) as f:
        info = yaml.safe_load(f)
        return Feast(slug=slug, **info)
